#!/usr/bin/env python3
"""生成验证报告的 Word 文档"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def create_docx_from_markdown():
    """从 Markdown 文件创建 Word 文档"""
    
    # 读取 Markdown 文件
    md_file = os.path.join(os.path.dirname(__file__), "验证报告.md")
    with open(md_file, "r", encoding="utf-8") as f:
        content = f.read()
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 解析内容
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # 处理标题
        if line.startswith('# '):
            # 一级标题
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.runs[0].font.size = Pt(18)
            heading.runs[0].bold = True
        elif line.startswith('## '):
            # 二级标题
            heading = doc.add_heading(line[3:], level=2)
            heading.runs[0].font.size = Pt(14)
            heading.runs[0].bold = True
        elif line.startswith('### '):
            # 三级标题
            heading = doc.add_heading(line[4:], level=3)
            heading.runs[0].font.size = Pt(12)
            heading.runs[0].bold = True
        elif line.startswith('|'):
            # 表格
            table_data = []
            # 读取表格行
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = [cell.strip() for cell in lines[i].strip().split('|')[1:-1]]
                if row and not all(c in ['-', ':', ' '] for c in ''.join(row)):
                    table_data.append(row)
                i += 1
            i -= 1  # 回退一行
            
            if table_data:
                # 创建表格
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_data
                        # 设置表头加粗
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
        elif line.startswith('- '):
            # 列表项
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('✅') or line.startswith('**'):
            # 强调文本
            p = doc.add_paragraph()
            run = p.add_run(line)
            if '✅' in line:
                run.bold = True
                run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
        else:
            # 普通段落
            p = doc.add_paragraph(line)
        
        i += 1
    
    # 保存文档
    output_file = os.path.join(os.path.dirname(__file__), "验证报告.docx")
    doc.save(output_file)
    print(f"✅ Word 文档已生成: {output_file}")

if __name__ == "__main__":
    try:
        create_docx_from_markdown()
    except ImportError:
        print("❌ 错误: 需要安装 python-docx 库")
        print("\n安装方法（选择其一）:")
        print("1. 使用虚拟环境:")
        print("   python3 -m venv venv")
        print("   source venv/bin/activate  # macOS/Linux")
        print("   pip install python-docx")
        print("\n2. 使用 pipx:")
        print("   brew install pipx")
        print("   pipx install python-docx")
        print("\n3. 使用 --break-system-packages (不推荐):")
        print("   pip3 install --break-system-packages python-docx")
    except Exception as e:
        print(f"❌ 错误: {e}")
        import traceback
        traceback.print_exc()

